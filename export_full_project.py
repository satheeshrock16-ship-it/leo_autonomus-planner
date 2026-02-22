import os
from docx import Document
from docx.shared import Pt

PROJECT_ROOT = "."
OUTPUT_FILE = "FULL_PROJECT_EXPORT.docx"

doc = Document()
doc.add_heading("FULL PROJECT EXPORT", level=1)

for root, dirs, files in os.walk(PROJECT_ROOT):
    for file in files:
        # Skip venv and git folders
        if ".git" in root or "venv" in root or "__pycache__" in root:
            continue

        file_path = os.path.join(root, file)

        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()

            doc.add_heading(file_path, level=2)

            paragraph = doc.add_paragraph()
            run = paragraph.add_run(content)
            run.font.name = "Courier New"
            run.font.size = Pt(8)

            doc.add_page_break()

        except:
            continue

doc.save(OUTPUT_FILE)
print("Export complete:", OUTPUT_FILE)
